import io
import tempfile
from typing import Iterable

from django.http import JsonResponse, FileResponse
from docx import Document
from openpyxl import load_workbook
from ..models import *


def create_template(template_file: Template) -> io.FileIO:
    fp = tempfile.NamedTemporaryFile('w+b', delete=False)
    fp.write(template_file.bytes)
    return fp


def create_doc(template_file: Template, template_data: Iterable[Descriptor], worker_data: Info) -> io.FileIO:
    file = create_template(template_file)
    doc = Document(file.name)

    for element in template_data:
        reference = str(element)
        data = getattr(worker_data, element.column) if element.key is None else worker_data.other_info[element.key]

        for paragraph in doc.paragraphs:
            if reference in paragraph.text:
                paragraph.text = paragraph.text.replace(f'{{{{{reference}}}}}', data)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if reference in paragraph.text:
                            paragraph.text = paragraph.text.replace(f'{{{{{reference}}}}}', data)

    doc.save(template_file.name)
    return file


def create_xlsx(template_file: Template, template_data: Iterable[Descriptor], worker_data: Info) -> io.FileIO:
    file = create_template(template_file)
    wb = load_workbook(filename=file.name)

    for element in template_data:
        table, cell = element.position.split(':')
        data = getattr(worker_data, element.column) if element.key is None else worker_data.other_info[element.key]
        reference = str(element)

        ws = wb[table]
        ws[cell] = ws[cell].value.replace(f'{{{{{reference}}}}}', data)

    wb.save(template_file.name)
    return file


def check_template_data(template_data: Iterable[Descriptor], worker_info_data: Info):
    missing_data = []

    for template_element in template_data:
        if template_element.key not in worker_info_data.other_info or not hasattr(worker_info_data,
                                                                                  template_element.column):
            missing_data.append(template_element)

    if missing_data:
        return JsonResponse({'message': 'Missing template data in database',
                             'data': DescriptorSerializer(missing_data, many=True).data}, status=206)
    return None


def use_doc_template(template: Template, template_data: Iterable[Descriptor],
                     worker_info_data: Info) -> JsonResponse | FileResponse:
    return FileResponse(create_doc(template, template_data, worker_info_data),
                        filename=f"{template.name}.{template.type}")


def use_xlsx_template(template: Template, template_data: Iterable[Descriptor],
                      worker_info_data: Info) -> JsonResponse | FileResponse:
    return FileResponse(create_xlsx(template, template_data, worker_info_data),
                        filename=f"{template.name}.{template.type}")


def use_template(template_id: int, worker_id: int | None = None) -> JsonResponse | FileResponse:
    template: Template = Template.objects.filter(id=template_id).first()

    if not worker_id:
        return FileResponse(create_template(template), filename=f"{template.name}.{template.type}")

    template_data = template.descriptor_set.all()
    worker_data: Worker = Worker.objects.filter(id=worker_id).first()
    worker_info_data: Info = worker_data.info

    check_template_data(template_data, worker_info_data)
    match template.type:
        case 'docx':
            return use_doc_template(template, template_data, worker_info_data)
        case 'xlsx':
            return use_xlsx_template(template, template_data, worker_info_data)
        case otherwise:
            return JsonResponse({'message': f'Incorrect file extension {template.type}'}, status=409)
